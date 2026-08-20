"""多格式文档解析（阶段⑤）：把 .md / .txt / .docx / .pdf 统一抽成纯文本。

RAG 只能吃「文字」，真实知识库文档却五花八门。这一步把各种格式「翻译」成
纯文本，喂给后面的切块 + 向量化。

设计：不只看扩展名，先用文件头「魔法字节」嗅探真实类型——
- PDF 以 %PDF 开头
- docx 本质是 zip 压缩包（PK\\x03\\x04 开头）
这样即使 Gradio 上传后临时文件丢了扩展名，也能正确解析。
"""
from pathlib import Path


def parse_file(path) -> str:
    """按真实文件类型解析，返回纯文本。"""
    p = Path(path)
    head = p.read_bytes()[:8]

    if head.startswith(b"%PDF"):
        return _parse_pdf(p)
    if head[:2] == b"PK":
        return _parse_docx(p)
    # 其余一律当纯文本（.md / .txt / .csv 等）
    return p.read_text(encoding="utf-8", errors="ignore")


def _parse_docx(p: Path) -> str:
    """Word：正文段落 + 表格（表格最易被忽略，却常藏着关键数据）。"""
    import docx

    doc = docx.Document(str(p))
    parts = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _parse_pdf(p: Path) -> str:
    """PDF：逐页抽文本（扫描版 PDF 无文字层时 extract_text 为空，属正常限制）。"""
    import pdfplumber

    pages = []
    with pdfplumber.open(str(p)) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            if text.strip():
                pages.append(text)
    return "\n\n".join(pages)
